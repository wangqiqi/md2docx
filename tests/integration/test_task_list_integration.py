"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
import pytest
test task list integration 测试
"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter

"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
任务列表转换集成测试
"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter


class TestTaskListIntegration:
    """任务列表转换集成测试"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter

    @pytest.fixture
    def base_converter(self):
        """创建基础转换器实例"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
        return BaseConverter(debug=False)

    def test_basic_task_list(self, base_converter):
        """测试基本任务列表转换"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
        md_text = """

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
- [ ] 未完成任务
- [x] 已完成任务
"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
        doc = base_converter.convert(md_text)

        # 验证结果
        # 应该有至少2个段落
        assert len(doc.paragraphs) >= 2

        # 验证段落内容 - 只验证存在性，不验证具体位置
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "未完成任务" in all_text
        assert "已完成任务" in all_text

    def test_nested_task_list(self, base_converter):
        """测试嵌套任务列表"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
        md_text = """

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
- [ ] 主任务1
  - [ ] 子任务1.1
  - [x] 子任务1.2
- [x] 主任务2
  - [x] 子任务2.1
"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
        doc = base_converter.convert(md_text)

        # 验证结果
        # 应该有至少2个段落（主任务）
        assert len(doc.paragraphs) >= 2

        # 验证段落内容 - 只验证主任务的存在性，不验证子任务
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "主任务1" in all_text
        assert "主任务2" in all_text

        # 验证至少有一个子任务
        assert "子任务" in all_text

    def test_mixed_list(self, base_converter):
        """测试混合列表（普通列表和任务列表）"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
        md_text = """

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
- 普通列表项1
- [ ] 任务列表项1
- 普通列表项2
- [x] 任务列表项2
"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
        doc = base_converter.convert(md_text)

        # 验证结果
        # 应该有至少4个段落
        assert len(doc.paragraphs) >= 4

        # 验证段落内容 - 只验证存在性，不验证具体位置
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "普通列表项1" in all_text
        assert "任务列表项1" in all_text
        assert "普通列表项2" in all_text
        assert "任务列表项2" in all_text

    def test_task_list_with_other_elements(self, base_converter):
        """测试任务列表与其他元素混合"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
        md_text = """

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
# 待办事项

- [ ] 完成报告
- [x] 发送邮件

---

> 注意：请在截止日期前完成所有任务
"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
        doc = base_converter.convert(md_text)

        # 验证结果
        # 应该至少有5个段落
        assert len(doc.paragraphs) >= 5

        # 验证段落内容 - 只验证存在性，不验证具体位置
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "待办事项" in all_text
        assert "完成报告" in all_text
        assert "发送邮件" in all_text
        assert "注意" in all_text
