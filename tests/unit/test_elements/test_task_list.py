"""
任务列表转换器单元测试
"""

from unittest.mock import MagicMock

import pytest
from docx import Document

from mddocx.converter.elements.task_list import TaskListConverter


def test_init():
    """测试初始化"""
    converter = TaskListConverter()
    assert converter is not None
    assert converter.document is None
    assert converter.base_converter is None

    # 测试带基础转换器的初始化
    base_converter = MagicMock()
    converter = TaskListConverter(base_converter)
    assert converter.base_converter == base_converter


def test_document_not_set():
    """测试文档未设置的情况"""
    converter = TaskListConverter()
    with pytest.raises(ValueError):
        converter.convert(MagicMock())


def test_convert_task_list():
    """测试任务列表转换"""
    # 创建转换器
    converter = TaskListConverter()
    converter.set_document(Document())

    # 创建模拟任务列表token
    task_token = MagicMock()
    task_token.type = "task_list_item"
    task_token.attrGet.return_value = "checked"

    # 模拟内容
    content_token = MagicMock()
    content_token.content = "任务内容"
    task_token.children = [content_token]

    # 转换任务列表项
    paragraph = converter.convert((task_token, task_token))

    # 验证结果
    assert paragraph is not None


def test_convert_task_list_debug_mode():
    """测试任务列表转换 - 调试模式"""
    # 创建转换器（启用调试）
    converter = TaskListConverter()
    converter.debug = True
    converter.set_document(Document())

    # 创建模拟任务列表token
    task_token = MagicMock()
    task_token.type = "task_list_item"

    content_token = MagicMock()
    content_token.content = "[x] 已完成任务"

    # 转换任务列表项
    paragraph = converter.convert((task_token, content_token))

    # 验证结果
    assert paragraph is not None


def test_add_checkbox_various_inputs():
    """测试_add_checkbox方法 - 各种输入"""
    # 创建转换器
    converter = TaskListConverter()
    converter.set_document(Document())

    paragraph = converter.document.add_paragraph()
    paragraph.add_run("测试任务")

    # 测试未勾选状态
    converter._add_checkbox(paragraph, is_checked=False)
    assert "× 测试任务" in paragraph.text

    # 重新创建段落测试已勾选状态
    paragraph2 = converter.document.add_paragraph()
    paragraph2.add_run("测试任务2")
    converter._add_checkbox(paragraph2, is_checked=True)
    assert "√ 测试任务2" in paragraph2.text


def test_task_list_with_base_converter():
    """测试任务列表转换器 - 带基础转换器"""
    from mddocx.converter.base import BaseConverter

    # 创建基础转换器
    base_converter = BaseConverter(debug=True)

    # 创建任务列表转换器
    converter = TaskListConverter(base_converter)
    converter.set_document(Document())

    # 验证基础转换器被正确设置
    assert converter.base_converter == base_converter
    assert converter.debug is True

    # 创建模拟任务列表token
    task_token = MagicMock()
    task_token.type = "task_list_item"

    content_token = MagicMock()
    content_token.content = "[x] 已完成任务"

    # 转换任务列表项
    paragraph = converter.convert((task_token, content_token))

    # 验证结果
    assert paragraph is not None
