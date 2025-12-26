"""
列表转换器单元测试
"""

from unittest.mock import MagicMock

import pytest
from docx import Document

from mddocx.converter.elements.list import ListConverter


def test_init():
    """测试初始化"""
    converter = ListConverter()
    assert converter is not None
    assert converter.document is None
    assert converter.base_converter is None

    # 测试带基础转换器的初始化
    base_converter = MagicMock()
    converter = ListConverter(base_converter)
    assert converter.base_converter == base_converter


def test_document_not_set():
    """测试文档未设置的情况"""
    converter = ListConverter()
    with pytest.raises(ValueError):
        converter.convert(MagicMock())


def test_convert_unordered_list():
    """测试无序列表转换"""
    # 创建转换器
    converter = ListConverter()
    converter.set_document(Document())

    # 创建模拟无序列表token
    ul_token = MagicMock()
    ul_token.type = "bullet_list_open"

    li_token = MagicMock()
    li_token.type = "list_item_open"

    text_token = MagicMock()
    text_token.type = "text"
    text_token.content = "列表项内容"

    # 设置token结构
    ul_token.children = [li_token]
    li_token.children = [text_token]

    # 转换列表
    paragraphs = converter.convert((ul_token, ul_token))

    # 验证结果
    assert paragraphs is not None


def test_convert_nested_ordered_list():
    """测试嵌套有序列表转换"""
    converter = ListConverter()
    converter.set_document(Document())

    ol1_token = MagicMock()
    ol1_token.type = "ordered_list_open"
    ol1_token.content = ""

    li1_token = MagicMock()
    li1_token.type = "list_item_open"

    text1_token = MagicMock()
    text1_token.type = "text"
    text1_token.content = "第一级项目"

    ol2_token = MagicMock()
    ol2_token.type = "ordered_list_open"
    ol2_token.content = "  "

    li2_token = MagicMock()
    li2_token.type = "list_item_open"

    text2_token = MagicMock()
    text2_token.type = "text"
    text2_token.content = "第二级项目"

    ol1_token.children = [li1_token]
    li1_token.children = [text1_token, ol2_token]
    ol2_token.children = [li2_token]
    li2_token.children = [text2_token]

    paragraphs = converter.convert((ol1_token, ol1_token))
    assert paragraphs is not None


def test_get_list_info_edge_cases():
    """测试_get_list_info方法的边界情况"""
    converter = ListConverter()

    # 测试无type属性的token
    token1 = MagicMock()
    del token1.type

    level, is_ordered = converter._get_list_info(token1)
    assert level == 1
    assert is_ordered is False

    # 测试无content属性的token
    token2 = MagicMock()
    token2.type = "bullet_list_open"
    del token2.content

    level, is_ordered = converter._get_list_info(token2)
    assert level == 1
    assert is_ordered is False

    # 测试有序列表
    token3 = MagicMock()
    token3.type = "ordered_list_open"
    token3.content = "    "

    level, is_ordered = converter._get_list_info(token3)
    assert level == 3
    assert is_ordered is True


def test_convert_ordered_list():
    """测试有序列表转换"""
    # 创建转换器
    converter = ListConverter()
    converter.set_document(Document())

    # 创建模拟有序列表token
    ol_token = MagicMock()
    ol_token.type = "ordered_list_open"

    li_token = MagicMock()
    li_token.type = "list_item_open"

    text_token = MagicMock()
    text_token.type = "text"
    text_token.content = "有序列表项"

    # 设置token结构
    ol_token.children = [li_token]
    li_token.children = [text_token]

    # 转换列表
    paragraphs = converter.convert((ol_token, ol_token))

    # 验证结果
    assert paragraphs is not None


def test_convert_nested_list():
    """测试嵌套列表转换"""
    # 创建转换器
    converter = ListConverter()
    converter.set_document(Document())

    # 创建模拟嵌套列表token
    ul_token = MagicMock()
    ul_token.type = "bullet_list_open"

    li1_token = MagicMock()
    li1_token.type = "list_item_open"

    text1_token = MagicMock()
    text1_token.type = "text"
    text1_token.content = "第一级列表项"

    ul2_token = MagicMock()
    ul2_token.type = "bullet_list_open"

    li2_token = MagicMock()
    li2_token.type = "list_item_open"

    text2_token = MagicMock()
    text2_token.type = "text"
    text2_token.content = "第二级列表项"

    # 设置token结构
    ul_token.children = [li1_token]
    li1_token.children = [text1_token, ul2_token]
    ul2_token.children = [li2_token]
    li2_token.children = [text2_token]

    # 转换列表
    paragraphs = converter.convert((ul_token, ul_token))

    # 验证结果
    assert paragraphs is not None
