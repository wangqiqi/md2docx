"""
图片转换器单元测试
"""

from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from mddocx.converter.elements.image import ImageConverter


def test_init():
    """测试初始化"""
    converter = ImageConverter()
    assert converter is not None
    assert converter.document is None

    # 测试带基础转换器的初始化
    base_converter = MagicMock()
    converter = ImageConverter(base_converter)
    assert converter.base_converter == base_converter


def test_document_not_set():
    """测试文档未设置的情况"""
    converter = ImageConverter()
    with pytest.raises(ValueError):
        converter.convert(MagicMock())


@patch("docx.text.run.Run.add_picture")
@patch("requests.get")
def test_convert_image(mock_get, mock_add_picture):
    """测试图片转换"""
    # 创建转换器
    converter = ImageConverter()
    converter.set_document(Document())

    # 模拟图片下载
    mock_response = MagicMock()
    mock_response.content = b"fake_image_data"
    mock_response.status_code = 200
    mock_get.return_value = mock_response

    # 创建模拟图片token
    image_token = MagicMock()
    image_token.type = "image"
    image_token.attrs = {"src": "https://example.com/image.png", "title": "测试图片"}
    image_token.content = "测试图片"

    # 转换图片
    paragraph = converter.convert((image_token, image_token))

    # 验证结果
    assert paragraph is not None
    mock_get.assert_called_once_with("https://example.com/image.png", timeout=10)


@patch("docx.text.run.Run.add_picture")
@patch("requests.get")
def test_convert_image_download_error(mock_get, mock_add_picture):
    """测试图片下载失败的情况"""
    # 创建转换器
    converter = ImageConverter()
    converter.set_document(Document())

    # 模拟下载失败
    mock_get.side_effect = Exception("Download failed")

    # 创建模拟图片token
    image_token = MagicMock()
    image_token.type = "image"
    image_token.attrs = {"src": "https://example.com/image.png", "title": "测试图片"}
    image_token.content = "测试图片"

    # 转换应该不会抛出异常，而是返回None或者空段落
    paragraph = converter.convert((image_token, image_token))
    assert paragraph is not None


def test_convert_image_no_attrs():
    """测试图片转换 - 无属性"""
    # 创建转换器
    converter = ImageConverter()
    converter.set_document(Document())

    # 创建模拟图片token（无attrs）
    image_token = MagicMock()
    image_token.type = "image"
    # 不设置attrs属性

    # 转换图片
    paragraph = converter.convert((image_token, image_token))

    # 验证结果 - 应该返回paragraph（即使图片加载失败）
    assert paragraph is not None

    """测试图片转换 - 无src属性"""
    # 创建转换器
    converter = ImageConverter()
    converter.set_document(Document())

    # 创建模拟图片token（无src）
    image_token = MagicMock()
    image_token.type = "image"
    image_token.attrs = {"title": "测试图片"}  # 有title但无src

    # 转换图片
    paragraph = converter.convert((image_token, image_token))

    # 验证结果 - 应该返回paragraph（即使无法获取图片数据）
    assert paragraph is not None


@patch("docx.text.run.Run.add_picture")
@patch("requests.get")
def test_convert_image_with_sizes(mock_get, mock_add_picture):
    """测试图片转换 - 带尺寸信息"""
    # 创建转换器
    converter = ImageConverter()
    converter.set_document(Document())

    # 模拟图片下载
    mock_response = MagicMock()
    mock_response.content = b"fake_image_data"
    mock_response.status_code = 200
    mock_get.return_value = mock_response

    # 创建模拟图片token（带尺寸）
    image_token = MagicMock()
    image_token.type = "image"
    image_token.attrs = {"src": "https://example.com/image.png", "title": "测试图片"}
    image_token.content = "图片 100x200"

    # 转换图片
    paragraph = converter.convert((image_token, image_token))

    # 验证结果
    assert paragraph is not None


@patch("docx.text.run.Run.add_picture")
@patch("requests.get")
def test_convert_image_debug_mode(mock_get, mock_add_picture):
    """测试图片转换 - 调试模式"""
    # 创建转换器（启用调试）
    converter = ImageConverter()
    converter.debug = True
    converter.set_document(Document())

    # 模拟图片下载
    mock_response = MagicMock()
    mock_response.content = b"fake_image_data"
    mock_response.status_code = 200
    mock_get.return_value = mock_response

    # 创建模拟图片token
    image_token = MagicMock()
    image_token.type = "image"
    image_token.attrs = {"src": "https://example.com/image.png"}

    # 转换图片
    paragraph = converter.convert((image_token, image_token))

    # 验证结果
    assert paragraph is not None


@patch("requests.get")
def test_get_image_data_cache(mock_get):
    """测试图片数据缓存"""
    # 创建转换器
    converter = ImageConverter()

    # 模拟图片下载
    mock_response = MagicMock()
    mock_response.content = b"fake_image_data"
    mock_response.status_code = 200
    mock_get.return_value = mock_response

    # 第一次获取
    result1 = converter._get_image_data("https://example.com/image.png")
    assert result1 is not None

    # 第二次获取（应该从缓存中获取）
    result2 = converter._get_image_data("https://example.com/image.png")
    assert result2 is not None

    # 验证只调用了一次网络请求
    assert mock_get.call_count == 1


def test_get_image_data_local_file():
    """测试获取本地图片文件"""
    import os
    import tempfile

    # 创建转换器
    converter = ImageConverter()

    # 创建临时图片文件
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
        temp_file.write(b"fake_png_data")
        temp_path = temp_file.name

    try:
        # 测试本地文件路径
        result = converter._get_image_data(temp_path)
        assert result is not None

    finally:
        os.unlink(temp_path)


def test_get_image_data_local_file_not_found():
    """测试本地图片文件不存在"""
    # 创建转换器
    converter = ImageConverter()

    # 测试不存在的文件
    result = converter._get_image_data("/nonexistent/path/image.png")
    assert result is None


@patch("docx.text.run.Run.add_picture")
def test_convert_image_in_paragraph(mock_add_picture):
    """测试在段落中转换图片"""
    # 创建转换器
    converter = ImageConverter()
    converter.set_document(Document())

    paragraph = converter.document.add_paragraph()

    # 创建模拟图片token
    image_token = MagicMock()
    image_token.type = "image"
    image_token.attrs = {"src": "https://example.com/image.png"}

    # 在段落中转换图片
    result = converter.convert_in_paragraph(paragraph, image_token)

    # 验证结果 - 方法不返回值
    assert result is None


def test_image_converter_with_base_converter():
    """测试图片转换器 - 带基础转换器"""
    from mddocx.converter.base import BaseConverter

    # 创建基础转换器
    base_converter = BaseConverter(debug=True)

    # 创建图片转换器
    converter = ImageConverter(base_converter)

    # 验证基础转换器被正确设置
    assert converter.base_converter == base_converter
