"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
import pytest
test text 测试
"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter

"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
文本转换器测试模块
"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter


def test_basic_paragraph():
    """测试基本段落转换"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
    md_text = "这是一个简单的段落。"

    converter = BaseConverter()
    converter.register_converter("paragraph", TextConverter())

    doc = converter.convert(md_text)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "这是一个简单的段落。"


def test_bold_text():
    """测试粗体文本"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
    md_text = "这是**粗体**文本。"

    converter = BaseConverter()
    converter.register_converter("paragraph", TextConverter())

    doc = converter.convert(md_text)
    runs = doc.paragraphs[0].runs

    assert runs[0].text == "这是"
    assert runs[1].text == "粗体" and runs[1].bold
    assert runs[2].text == "文本。"


def test_italic_text():
    """测试斜体文本"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
    md_text = "这是*斜体*文本。"

    converter = BaseConverter()
    converter.register_converter("paragraph", TextConverter())

    doc = converter.convert(md_text)
    runs = doc.paragraphs[0].runs

    assert runs[0].text == "这是"
    assert runs[1].text == "斜体" and runs[1].italic
    assert runs[2].text == "文本。"


def test_strikethrough_text():
    """测试删除线文本"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
    md_text = "这是~~删除线~~文本。"

    converter = BaseConverter()
    converter.register_converter("paragraph", TextConverter())

    doc = converter.convert(md_text)
    runs = doc.paragraphs[0].runs

    assert runs[0].text == "这是"
    assert runs[1].text == "删除线" and runs[1].font.strike
    assert runs[2].text == "文本。"


def test_multiple_paragraphs():
    """测试多个段落"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
    md_text = """第一段。

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter

第二段。

第三段。"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter

    converter = BaseConverter()
    converter.register_converter("paragraph", TextConverter())

    doc = converter.convert(md_text)
    paragraphs = doc.paragraphs

    assert len(paragraphs) == 3
    assert paragraphs[0].text == "第一段。"
    assert paragraphs[1].text == "第二段。"
    assert paragraphs[2].text == "第三段。"


def test_empty_paragraph():
    """测试空段落处理"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
    converter = TextConverter()
    converter.set_document(Document())

    md = MarkdownIt()
    tokens = md.parse("")  # 空文本

    # 应该不会抛出异常
    if len(tokens) >= 2:
        converter.convert((tokens[0], tokens[1]))


def test_inline_code():
    """测试行内代码"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
    md_text = "这是`行内代码`的示例。"

    converter = BaseConverter()
    converter.register_converter("paragraph", TextConverter())

    doc = converter.convert(md_text)
    runs = doc.paragraphs[0].runs

    assert runs[0].text == "这是"
    # 行内代码应该直接显示代码内容，不带反引号
    assert runs[1].text == "行内代码"
    # 检查是否使用了等宽字体
    assert runs[1].font.name == "Consolas"
    assert runs[2].text == "的示例。"


def test_document_not_set():
    """测试未设置文档实例的情况"""

import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt

from mddocx.converter.base import BaseConverter
from mddocx.converter.elements.hr import HRConverter
from mddocx.converter.elements.html import HtmlConverter
from mddocx.converter.elements.image import ImageConverter
from mddocx.converter.elements.table import TableConverter
from mddocx.converter.elements.task_list import TaskListConverter
from mddocx.converter.elements.text import TextConverter
    converter = TextConverter()
    md = MarkdownIt()
    tokens = md.parse("测试文本")

    with pytest.raises(ValueError, match="Document not set"):
        converter.convert((tokens[0], tokens[1]))
