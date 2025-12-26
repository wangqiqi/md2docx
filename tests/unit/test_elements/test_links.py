"""
链接转换器单元测试
"""

from unittest.mock import MagicMock

import pytest
from docx import Document

from mddocx.converter.elements.links import LinkConverter


class TestLinkConverter:
    """链接转换器测试类"""

    @pytest.fixture
    def converter(self):
        """创建链接转换器实例"""
        conv = LinkConverter()
        conv.set_document(Document())
        return conv

    def test_init(self):
        """测试初始化"""
        converter = LinkConverter()
        assert converter is not None
        assert converter.document is None

        # 测试带基础转换器的初始化
        base_converter = MagicMock()
        converter = LinkConverter(base_converter)
        assert converter.base_converter == base_converter

    def test_document_not_set(self):
        """测试文档未设置的情况"""
        converter = LinkConverter()
        with pytest.raises(ValueError):
            converter.convert(MagicMock())

    def test_set_document_none(self):
        """测试设置空文档"""
        converter = LinkConverter()
        with pytest.raises(ValueError):
            converter.set_document(None)

    def test_convert_link(self, converter):
        """测试链接转换"""
        # 创建模拟链接token
        link_token = MagicMock()
        link_token.attrs = {"href": "https://example.com"}

        content_token = MagicMock()
        content_token.content = "链接文本"

        # 转换链接
        result = converter.convert((link_token, content_token))

        # 验证结果 - convert方法不返回任何值
        assert result is None

        # 验证文档中添加了段落
        assert len(converter.document.paragraphs) >= 1
        paragraph = converter.document.paragraphs[-1]
        assert "链接文本" in paragraph.text

    # def test_convert_link_empty_attrs(self, converter):
    #     """测试链接转换 - 空属性"""
    #     # 创建模拟链接token（无attrs）
    #     link_token = MagicMock()
    #     del link_token.attrs  # 确保没有attrs属性

    #     content_token = MagicMock()
    #     content_token.content = "链接文本"

    #     # 转换链接
    #     result = converter.convert((link_token, content_token))

    #     # 验证结果 - 应该返回None
    #     assert result is None

    #     # 文档应该没有变化
    #     assert len(converter.document.paragraphs) == 0

    # def test_convert_link_url_only(self, converter):
    #     """测试链接转换 - 只有URL"""
    #     # 创建模拟链接token
    #     link_token = MagicMock()
    #     link_token.attrs = {"href": "https://example.com"}

    #     content_token = MagicMock()
    #     content_token.content = ""  # 空文本

    #     # 转换链接
    #     result = converter.convert((link_token, content_token))

    #     # 验证结果
    #     assert result is None

    #     # 验证文档中添加了段落，文本使用URL
    #     assert len(converter.document.paragraphs) >= 1
    #     paragraph = converter.document.paragraphs[-1]
    #     assert "https://example.com" in paragraph.text

    # def test_convert_link_empty_content(self, converter):
    #     """测试链接转换 - 空内容"""
    #     # 创建模拟链接token
    #     link_token = MagicMock()
    #     link_token.attrs = {"href": "https://example.com"}

    #     content_token = MagicMock()
    #     del content_token.content  # 确保没有content属性

    #     # 转换链接
    #     result = converter.convert((link_token, content_token))

    #     # 验证结果
    #     assert result is None

    #     # 验证文档中添加了段落，文本使用URL
    #     assert len(converter.document.paragraphs) >= 1
    #     paragraph = converter.document.paragraphs[-1]
    #     assert "https://example.com" in paragraph.text

    def test_convert_link_no_url(self, converter):
        """测试链接转换 - 无URL"""
        # 创建模拟链接token
        link_token = MagicMock()
        link_token.attrs = {}  # 无href

        content_token = MagicMock()
        content_token.content = "链接文本"

        # 转换链接
        result = converter.convert((link_token, content_token))

        # 验证结果
        assert result is None

        # 验证文档中添加了段落
        assert len(converter.document.paragraphs) >= 1
        paragraph = converter.document.paragraphs[-1]
        assert "链接文本" in paragraph.text

    def test_convert_link_empty_text_and_url(self, converter):
        """测试链接转换 - 空文本和URL"""
        # 创建模拟链接token
        link_token = MagicMock()
        link_token.attrs = {"href": ""}

        content_token = MagicMock()
        content_token.content = ""

        # 转换链接
        result = converter.convert((link_token, content_token))

        # 验证结果
        assert result is None

        # 验证文档中添加了段落，显示"(空链接)"
        assert len(converter.document.paragraphs) >= 1
        paragraph = converter.document.paragraphs[-1]
        assert "(空链接)" in paragraph.text

    def test_convert_in_paragraph_with_children(self, converter):
        """测试在段落中转换链接 - 有子元素"""
        paragraph = converter.document.add_paragraph()

        # 创建模拟链接token
        link_token = MagicMock()
        link_token.type = "link"
        link_token.attrs = {"href": "https://example.com"}
        link_token.children = [
            MagicMock(type="text", content="链"),
            MagicMock(type="text", content="接"),
            MagicMock(type="text", content="文"),
            MagicMock(type="text", content="本"),
        ]

        # 转换链接
        result = converter.convert_in_paragraph(paragraph, link_token)

        # 验证结果 - 方法不返回值
        assert result is None

        # 验证段落中有链接文本
        assert "链接文本" in paragraph.text

    def test_convert_in_paragraph_link_text_provided(self, converter):
        """测试在段落中转换链接 - 提供链接文本"""
        paragraph = converter.document.add_paragraph()

        # 创建模拟链接token
        link_token = MagicMock()
        link_token.type = "link"
        link_token.attrs = {"href": "https://example.com"}

        # 转换链接，提供自定义文本
        result = converter.convert_in_paragraph(
            paragraph, link_token, link_text="自定义链接"
        )

        # 验证结果
        assert result is None
        assert "自定义链接" in paragraph.text

    def test_convert_in_paragraph_no_attrs(self, converter):
        """测试在段落中转换链接 - 无属性"""
        paragraph = converter.document.add_paragraph()

        # 创建模拟链接token（无attrs）
        link_token = MagicMock()
        link_token.type = "link"

        # 转换链接
        result = converter.convert_in_paragraph(paragraph, link_token)

        # 验证结果 - 应该返回None
        assert result is None

        # 段落应该没有变化
        assert paragraph.text == ""

    def test_convert_in_paragraph_empty_attrs(self, converter):
        """测试在段落中转换链接 - 空属性"""
        paragraph = converter.document.add_paragraph()

        # 创建模拟链接token（空attrs）
        link_token = MagicMock()
        link_token.type = "link"
        link_token.attrs = {}

        # 转换链接
        result = converter.convert_in_paragraph(paragraph, link_token)

        # 验证结果
        assert result is None
        assert paragraph.text == ""

    def test_convert_in_paragraph_with_style(self, converter):
        """测试在段落中转换链接 - 带样式"""
        paragraph = converter.document.add_paragraph()

        # 创建模拟链接token
        link_token = MagicMock()
        link_token.type = "link"
        link_token.attrs = {"href": "https://example.com"}
        link_token.children = [MagicMock(type="text", content="样式链接")]

        # 转换链接，带样式
        style = {"bold": True, "italic": True}
        result = converter.convert_in_paragraph(paragraph, link_token, style=style)

        # 验证结果
        assert result is None
        assert "样式链接" in paragraph.text

    def test_ensure_hyperlink_style(self, converter):
        """测试确保超链接样式存在"""
        # 样式应该在set_document时创建
        assert "Hyperlink" in converter.document.styles

        # 再次调用不应该出错
        converter._ensure_hyperlink_style()
        assert "Hyperlink" in converter.document.styles
