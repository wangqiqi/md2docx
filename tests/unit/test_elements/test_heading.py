"""
标题转换器测试模块
"""

import pytest
from docx import Document

from mddocx.converter.elements.heading import HeadingConverter


class TestHeadingConverter:
    """标题转换器测试类"""

    @pytest.fixture
    def converter(self):
        """创建标题转换器实例"""
        conv = HeadingConverter()
        conv.document = Document()
        return conv

    def test_heading_conversion(self, converter):
        """测试标题转换功能"""
        # 准备测试数据 - 模拟 markdown-it 的 token 结构
        heading_token = type("Token", (), {"tag": "h1", "type": "heading_open"})()
        content_token = type("Token", (), {"content": "一级标题", "type": "text"})()

        # 执行转换
        converter.convert((heading_token, content_token))

        # 验证结果
        assert len(converter.document.paragraphs) == 1
        paragraph = converter.document.paragraphs[0]
        assert paragraph.text == "一级标题"
        assert paragraph.style.name == "Heading 1"

    def test_invalid_heading_level(self, converter):
        """测试无效标题级别"""
        # 测试无效的标题级别
        heading_token = type("Token", (), {"tag": "h7", "type": "heading_open"})()
        content_token = type("Token", (), {"content": "无效标题", "type": "text"})()

        # 这应该抛出异常
        with pytest.raises(ValueError, match="不支持的标题级别"):
            converter.convert((heading_token, content_token))

    def test_document_not_set(self):
        """测试文档未设置的情况"""
        converter = HeadingConverter()
        # 不设置 document 属性

        heading_token = type("Token", (), {"tag": "h1", "type": "heading_open"})()
        content_token = type("Token", (), {"content": "测试标题", "type": "text"})()

        # 这应该抛出异常
        with pytest.raises(ValueError, match="Document not set"):
            converter.convert((heading_token, content_token))
