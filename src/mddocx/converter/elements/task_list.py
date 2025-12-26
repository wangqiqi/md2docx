"""
任务列表转换器模块
"""

import re

from .base import ElementConverter
from .list import ListConverter


class TaskListConverter(ElementConverter):
    """任务列表转换器，处理Markdown中的任务列表（TODO列表）"""

    def __init__(self, base_converter=None):
        """初始化任务列表转换器

        Args:
            base_converter: 基础转换器实例
        """
        super().__init__(base_converter)
        self.debug = False
        self.list_converter = None
        if base_converter:
            self.debug = base_converter.debug
            # 获取列表转换器，用于处理基本列表结构
            if "list" in base_converter.converters:
                self.list_converter = base_converter.converters["list"]
            else:
                self.list_converter = ListConverter(base_converter)

    def convert(self, tokens):
        """转换任务列表token为DOCX带符号的列表

        Args:
            tokens: 任务列表token元组 (list_token, content_token)

        Returns:
            docx.paragraph: 创建的段落对象
        """
        if not self.document:
            raise ValueError("Document not set for TaskListConverter")

        if self.debug:
            print(f"处理任务列表: {tokens}")

        # 解析token
        list_token, content_token = tokens

        # 检查是否为任务列表项
        is_checked = False
        task_text = ""

        # 清理任务标记的正则表达式
        task_pattern = re.compile(r"^\s*\[([ x])\]\s*")

        if hasattr(content_token, "content"):
            content = content_token.content
            # 使用正则表达式检查和移除任务标记
            try:
                match = task_pattern.match(content)
                if match:
                    is_checked = match.group(1) == "x"
                    task_text = task_pattern.sub("", content)
                else:
                    task_text = content
            except TypeError:
                # 处理 content 不是字符串的情况
                task_text = str(content) if content is not None else ""
        elif hasattr(content_token, "children"):
            for child in content_token.children:
                if hasattr(child, "type") and child.type == "checkbox_input":
                    is_checked = (
                        child.attrs.get("checked", False)
                        if hasattr(child, "attrs")
                        else False
                    )
                elif hasattr(child, "content"):
                    # 使用正则表达式移除内容中的任务标记
                    try:
                        child_content = child.content
                        match = task_pattern.match(child_content)
                        if match:
                            child_content = task_pattern.sub("", child_content)
                        task_text += child_content
                    except TypeError:
                        # 处理 child.content 不是字符串的情况
                        task_text += (
                            str(child.content) if child.content is not None else ""
                        )

        # 使用符号替代复选框
        checkbox_symbol = "☐ " if not is_checked else "☑ "

        # 将符号添加到任务文本前面
        task_text_with_symbol = checkbox_symbol + task_text

        # 创建普通段落，不使用列表样式，避免重复的列表符号
        paragraph = self.document.add_paragraph()
        paragraph.add_run(task_text_with_symbol)

        # 添加适当的缩进，让任务列表看起来像列表项
        # 获取列表层级（从 list_token 中推断）
        level = 1
        if hasattr(list_token, "content"):
            indent = len(list_token.content)
            level = (indent // 2) + 1

        # 设置缩进
        indent_inches = 0.25 * (level - 1)
        from docx.shared import Inches

        paragraph.paragraph_format.left_indent = Inches(indent_inches)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)

        return paragraph

    def _add_checkbox(self, paragraph, is_checked=False):
        """向段落添加复选框

        Args:
            paragraph: 段落对象
            is_checked: 是否勾选
        """
        # 检查段落是否为None
        if paragraph is None:
            if self.debug:
                print("警告: 尝试向None段落添加复选框")
            return

        # 获取段落的第一个run
        if not paragraph.runs:
            run = paragraph.add_run()
        else:
            run = paragraph.runs[0]

        # 添加复选框符号
        if is_checked:
            run.text = "√ " + run.text
        else:
            run.text = "× " + run.text
